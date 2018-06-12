import nltk
from nltk.corpus import cess_esp as cess
from nltk import word_tokenize
from nltk import sent_tokenize
from docx import Document
from nltk import UnigramTagger as ut
from nltk import BigramTagger as bt

# Read the corpus into a list, 
# each entry in the list is one sentence.
cess_sents = cess.tagged_sents()
# Train the unigram tagger
uni_tag = ut(cess_sents)

text = ""
f = open('PLIEGO_EJEMPLO.docx', 'rb')
document = Document(f)
for i in document.paragraphs:
    text +=i.text
f.close()

print(text[:45])

# obtengo los tokens
tokens = word_tokenize(text)
words = tokens[:50]
# Tagger reads a list of tokens.
resUni = uni_tag.tag(words)
print(resUni)
# Split corpus into training and testing set.
train = int(len(cess_sents)*90/100) # 90%
print(train)
# Train a bigram tagger with only training data.
bi_tag = bt(cess_sents[:train])

# Evaluates on testing data remaining 10%
bi_tag.evaluate(cess_sents[train+1:])

# Using the tagger.
res = bi_tag.tag(words)

print(res)
#from io import StringIO
#nltk.data.load('tokenizers/punkt/spanish.pickle') 
'''
text = ""
f = open('PLIEGO_EJEMPLO.docx', 'rb')
document = Document(f)
for i in document.paragraphs:
    text +=i.text
f.close()

print(text[:45])

# obtengo los tokens
tokens = word_tokenize(text)
print(type(tokens))
print(len(tokens))
print(tokens[:10])

words = tokens[:10]
tagged_words = nltk.corpus.cess_esp.pos_tag(words)
pos_tag
nouns = []
for (word, tag) in tagged_words:
    print(word+' '+tag)
'''